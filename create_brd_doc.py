from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import json

def add_bordered_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.space_before = Inches(0)  # No space before section headers
    run = paragraph.add_run(text)
    font = run.font
    font.size = Pt(16)  # Section headers are 16pt
    font.bold = True
    
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    
    # Add light blue background
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'E6F3FF')
    pPr.append(shd)
    
    # Add borders
    pBdr = OxmlElement('w:pBdr')
    borders = [('w:top', 6), ('w:left', 6), ('w:bottom', 6), ('w:right', 6)]
    for border, size in borders:
        borderElement = OxmlElement(border)
        borderElement.set(qn('w:val'), 'single')
        borderElement.set(qn('w:sz'), str(size))
        borderElement.set(qn('w:color'), '0000FF')
        pBdr.append(borderElement)
    
    pPr.append(pBdr)
    return paragraph

def add_subsection_with_content(doc, subsection_name, content):
    # Add subsection heading with minimal spacing
    subsection = doc.add_paragraph()
    subsection.space_before = Inches(0.1)  # Small space before subsection
    run = subsection.add_run(subsection_name)
    run.font.size = Pt(14)  # Subsection headers are 14pt
    run.font.bold = True
    
    # Add content with single line spacing
    content_para = doc.add_paragraph()
    content_para.space_before = Inches(0)
    content_run = content_para.add_run(content)
    content_run.font.size = Pt(14)  # Content is 14pt
    
    # Add exactly one line space after content
    content_para.space_after = Inches(0.2)

def create_brd_document(json_file_path, output_file_path):
    with open(json_file_path, 'r') as file:
        data = json.load(file)
    
    doc = Document()
    
    # Title with 18pt and underline
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Business Requirement Documentation")
    title_run.font.size = Pt(18)
    title_run.font.underline = WD_UNDERLINE.SINGLE
    title_run.font.bold = True
    title.space_after = Inches(0)  # No space after title
    
    # Executive Summary
    add_bordered_heading(doc, "Executive Summary")
    para = doc.add_paragraph()
    para.space_before = Inches(0)
    para.add_run(data["Executive Summary"]).font.size = Pt(14)
    para.space_after = Inches(0.2)
    
    # Project Objectives
    add_bordered_heading(doc, "Project Objectives")
    for key, value in data["Project Objectives"].items():
        add_subsection_with_content(doc, f"{key}:", value)
    
    # Project Scope
    add_bordered_heading(doc, "Project Scope")
    for key, value in data["Project Scope"].items():
        if isinstance(value, list):
            bullet_list = "\n".join(f"• {item}" for item in value)
            add_subsection_with_content(doc, f"{key}:", bullet_list)
        else:
            add_subsection_with_content(doc, f"{key}:", value)
    
    # Business Requirements
    add_bordered_heading(doc, "Business Requirements")
    para = doc.add_paragraph()
    para.space_before = Inches(0)
    para.add_run(data["Business Requirements"]).font.size = Pt(14)
    para.space_after = Inches(0.2)
    
    # Key Stakeholders
    add_bordered_heading(doc, "Key Stakeholders")
    bullet_list = "\n".join(f"• {item}" for item in data["Key Stakeholders"])
    para = doc.add_paragraph()
    para.space_before = Inches(0)
    para.add_run(bullet_list).font.size = Pt(14)
    para.space_after = Inches(0.2)
    
    # Project Constraints
    add_bordered_heading(doc, "Project Constraints")
    for key, value in data["Project Constraints"].items():
        add_subsection_with_content(doc, f"{key}:", value)
    
    # Cost-Benefit Analysis
    add_bordered_heading(doc, "Cost-Benefit Analysis")
    bullet_list = "\n".join(f"• {item}" for item in data["Cost-Benefit Analysis"])
    para = doc.add_paragraph()
    para.space_before = Inches(0)
    para.add_run(bullet_list).font.size = Pt(14)
    
    doc.save(output_file_path)