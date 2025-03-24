#This file contains all the agent functionalities

#Importing all the necessaries
import template
import variables
import regex
import ast
import pathlib
import fitz
import random
import os
from langchain_core.messages import HumanMessage, SystemMessage
from langchain.text_splitter import CharacterTextSplitter
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langchain_core.messages import AIMessage
#setting upt the API keys
os.environ['OPENAI_API_KEY'] = ''
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = ""

#Setting up the model
model = ChatOpenAI(model='gpt-3.5-turbo', temperature = 0)

end_convo = 0
# A function to generate random numbers. This helps in generating different responses providing the same meaning to randomize the response lines.
def gen_random(i,j):
    random_number = random.randint(i,j)-1
    return random_number

#Function used to communicate with the model.
def ask_gpt(system_message, human_message):
    messages = [
        SystemMessage(system_message),
        HumanMessage(human_message),
    ]
    return model.invoke(messages).content

#Function to greet the user
def greeter():
    time = datetime.now().hour  # Get the current hour in 24-hour format
    if 5 <= time < 12:
        greetings = 'Good Morning'
    elif 12 <= time < 18:
        greetings = 'Good Afternoon'
    elif 18 <= time <= 23:
        greetings = 'Good Evening'
    else:
        greetings = 'Hard Worker'
    return greetings

#Function to update the user choice of providing information. As of now as the model is only generated for conversation based inputs, it is returning 0.
def update_user_choice(choice_input):
    if choice_input == 1:
        return 0
    else:
        return 1

def option_2(path):
    ext = pathlib.Path(path).suffix
    if ext == '.pdf':
        texts = []
        new_text = ''
        doc = fitz.open(path)
        for i in range(doc.page_count):
            p = doc.load_page(i)
            texts.append(p.get_text())
        for i in range(len(texts)):
            new_text += ''.join(texts[i])
        return new_text

    if ext == '.docx':
        texts = []
        new_text = ''
        doc = Document(path)
        for para in doc.paragraphs:
            texts.append(para.text)
        for i in range(len(texts)):
            new_text += ''.join(texts[i])
        new_text = new_text.replace('\n', '')
        return new_text

    if ext == '.txt':
        texts = []
        new_text = ''
        with open(path, 'r') as file:
            text = file.read()
            texts.append(text)
        for i in range(len(texts)):
            new_text += ''.join(texts[i])
        return new_text

#Function to record the responses of the user related to project in a text document store at the backend
def record_response(response):
    with open(variables.project_file,'a', encoding='utf-8') as f:
        f.write(response+'\n')

#Function to spit the text in the text document where responses related to project is stored (Currently all uses of this function are reduced).
def text_splitter():
    raw_text=''
    with open(variables.project_file, 'r') as f:
        for i in f.readlines():
            raw_text+=i
        raw_text = raw_text.replace('. ', '.\n')
    splitter = CharacterTextSplitter(
        separator = '.\n',
        chunk_size = 800,
        chunk_overlap = 200,
        length_function = len
    )
    texts = splitter.split_text(raw_text)
    return texts

#A part of improved requirement loop, with functionality of analysing text and returns the missing fields and subfields.
def text_analysis(structure):
    missing_fields = {}

    raw_text=''
    with open(variables.project_file, 'r') as f:
        for i in f.readlines():
            raw_text+=i
        text = raw_text.replace('. ', '.\n')
    llm = ChatOpenAI(model="gpt-3.5-turbo")

    # The prompt setup
    prompt = PromptTemplate(
        input_variables=["context", "question"],
        template="You are a text analysis expert.\nGiven the following context:\n{context}\n"
                 "Answer as per the instructions mentioned in the question:\n{question}",
    )
    chain = prompt | llm

    print("Analyzing 🧐.....")

    for i in structure:
        missing_subfield = []
        for j in structure[i]:

            # Query setup
            query = f'''
                Tasks:
                1. The main task is to analyze the complete given text and find the information related to the field '{j}'.
                2. If any sentence is related to the field '{j}', return the answer as 'yes'.
                3. If no or very little information is related to the field '{j}' in the given text, return the answer as 'no'.

                Examples:
                - If the field is 'Problem statement' and the context is "The project is called 'Tizum'. Goals and objectives are to reduce carbon footprint.", the answer is 'no'(context doesn't contains field 'Problem statement').
                - If the field is 'Problem statement' and the context is "The project is called 'Tizum'. Goals and objectives are to reduce carbon footprint. However, technological growth contributes to pollution, which harms the environment.", the answer is 'yes'(context contains field 'Problem statement').
                '''
            #Use chain.invoke
            result = chain.invoke({"context": text, "question": query})

            if isinstance(result, AIMessage):
                processed_result = result.content.strip().lower()  # Extract text properly
            else:
                raise TypeError(f"Unexpected output format: {type(result)}")

            if processed_result == 'no':
                missing_subfield.append(j)

        if missing_subfield:
            missing_fields.update({i: missing_subfield})

    return missing_fields

#A part of improved requirement loop, which asks question to user, make conversations and records the response based on the response of the user.
def response_analysis(missing_fields):
    fields = missing_fields.copy()
    global end_convo
    if len(missing_fields)>0:
        for i in fields:
            for j in (fields[i])[:]:
                ran_number = gen_random(1, len(variables.question_lines_1))
                request = f"{variables.question_lines_1[ran_number]}{i}{variables.question_lines_2[ran_number]}{j}?"
                print(request)
                response = input()
                print('--------------------------------------------------------------')
                variables.request_response.update({request: response})
                # Sentiment analysis of the user
                sys_msg = f'''
                Analyze the given input and categorize the response based on its relevance and alignment with {j}:

                1. Irrelevant: If the input provides information that is unrelated to {j} or lacks relevance, respond with 'irrelevant'.

                2. Relevant: If the input explicitly provides a valid answer to {j} (e.g., direct or specific responses such as dates, names, values, details, or related information), respond with 'relevant'.

                3. No: If the input conveys an unwillingness to provide an answer or expresses a sentiment of refusal, respond with 'no'.

                4. Present: If the input suggests that the requested information has already been provided (e.g., the information is mentioned before, answered earlier, or referred to as already available), respond with 'present'.

                5. End Requirement: If the input provides sufficient information to indicate task completeness or includes a request or command for further actions (e.g., generating content based on the provided information), respond with 'end requirement'.

                Carefully analyze whether the input explicitly or implicitly answers {j}, ensuring that direct and concise answers such as specific terms, dates, or values are not overlooked.'''

                # Agent responses with proper messages
                variables.user_responses.append(ask_gpt(sys_msg, response))
                if variables.user_responses[-1].strip().lower() == 'no':
                    print('Okay. No issues 👍...')
                    missing_fields[i].remove(j)
                    record_response(f'No responses on the {i}------{j}')
                elif variables.user_responses[-1].strip().lower() == 'relevant':
                    print("Great 😊...")
                    record_response(f"{j}: {response}")
                elif variables.user_responses[-1].strip().lower() == 'irrelevant':
                    print("The answer is probably irrelevant to the asked field 🤔....")
                elif variables.user_responses[-1].strip().lower() == 'present':
                    print("Okay. Sorry for the inconvenience 🥺...")
                elif variables.user_responses[-1].strip().lower() == 'end requirement':
                    print('Sure. Here you go...')
                    return {},1

            if not missing_fields[i]:
                missing_fields.pop(i)
    if len(missing_fields)>0:
        text_analysis(missing_fields)
    else:
        end_convo = 1
    return missing_fields, end_convo

#A function which extracts the information for each fields from the given details.
def text_extraction():
    raw_text = ''
    with open(variables.project_file, 'r') as f:
        for i in f.readlines():
            raw_text += i
        text = raw_text.replace('. ', '.\n')
    llm = ChatOpenAI(model="gpt-3.5-turbo")

    # The prompt setup
    prompt = PromptTemplate(
        input_variables=["context", "question"],
        template="You are a text analysis and extraction expert.\nGiven the following context:\n{context}\n"
                 "Answer as per the instructions mentioned in the question:\n{question}",
    )
    chain = prompt | llm
    sys_msg = '''
    If the given input is not in an 'expected format', return in the expected format.
    '''
    print('Information is being extracted ...')
    sections = list(template.template_for_charter.keys())
    for i in sections:
        processed_output = {}
        for j in template.template_for_charter[i]:
            if i == 'Project Overview':
                if j == 'Project Name':
                    query = f'''
                    The description of the subfield{j}: The official title or identifier of the project for easy reference.
                    If the input text contains '{j}' or related meanings, sentiments, or similar words, extract and return the relevant processed content. Skip lines where information is missing or indicates it cannot be provided. Exclude any mention of the phrase 'no responses on the {i}------{j}' or similar. Only return the extracted information without additional statements, explanations, or formatting.
                    For example:
                    If input is 'The name of the project is xyz.', then return 'xyz'; dont return '{j}:The name of the project is xyz' or something which is unnecessary. If the input contains some something which means the information of {j} cannot be provided as mentioned above, then return nothing.
                    '''
                    
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})

                if j == 'Project Start(projected)':
                    query = f'''
                    The description of the subfield{j}: The anticipated date when the project activities will begin.
                    If the input text contains '{j}' or related meanings, sentiments, or similar words, extract and return the relevant processed content. Skip lines where information is missing or indicates it cannot be provided. Exclude any mention of the phrase 'no responses on the {i}------{j}' or similar. Only return the extracted information without additional statements, explanations, or formatting.
                    For example:
                    If the input has 'The start date of the project is dd/mm/yyyy' or in any other format, then return in month date, year format. Try to add suffix to the date like 'st', 'nd', 'th' 'rd'. No need of adding anything in addition to the actual date like, dont return 'The start date is month date, year' or any any unnecessary date.
                    '''
                    
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})

                if j == 'Project End(projected)':
                    query = f'''
                    The description of the subfield{j}: The estimated completion date for the project.
                    If the input text contains '{j}' or related meanings, sentiments, or similar words, extract and return the relevant processed content. Skip lines where information is missing or indicates it cannot be provided. Exclude any mention of the phrase 'no responses on the {i}------{j}' or similar. Only return the extracted information without additional statements, explanations, or formatting.
                    For example:
                    If the input has 'The start date of the project is dd/mm/yyyy' or in any other format, then return in month date, year format. Try to add suffix to the date like 'st', 'nd', 'th', 'rd'. No need of adding anything in addition to the actual date like, dont return 'The start date is month date, year' or any any unnecessary date.
                    '''
                    
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})

                if j == 'Authorization Date':
                    query = f'''
                    The description of the subfield{j}: The date when the project received formal approval to start.
                    If the input text contains '{j}' or related meanings, sentiments, or similar words, extract and return the relevant processed content. Skip lines where information is missing or indicates it cannot be provided. Exclude any mention of the phrase 'no responses on the {i}------{j}' or similar. Only return the extracted information without additional statements, explanations, or formatting.
                    For example:
                    If the input has 'The start date of the project is dd/mm/yyyy' or in any other format, then return in month date, year format. Try to add suffix to the date like 'st', 'nd', 'th' 'rd'. No need of adding anything in addition to the actual date like, dont return 'The start date is month date, year' or any any unnecessary date.
                    '''
                    
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})

            if i == 'Business Case and Project Objectives':
                if j == 'Problem Statement':
                    query = f"""
                    The description of the subfield {j}: A concise description of the issue, challenge, or opportunity that the project aims to address.

                    Instructions:
                    1. Extract and return only the content directly associated with the problem statement.
                    2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                    3. Do not include any prefixes, headers, or phrases like 'the problem is,' 'problem statement:', or similar.

                    Note:
                    - Return only the extracted problem statement associated with '{j}', in its original form, without rephrasing or introducing additional commentary.
                    - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                    Example:
                    Suppose the given text is 'The problem is that the world population is high and creating unnecessary fights for every resource, leading to unhealthy competition.'
                    Then return: 'The world population is high and creating unnecessary fights for every resource, leading to unhealthy competition.'

                    Suppose the given text is 'Our main challenge is the increasing demand for clean water in urban areas, causing supply shortages.'
                    Then return: 'The increasing demand for clean water in urban areas, causing supply shortages.'

                    Go through the instructions and examples carefully to structure the response as requested.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})
                if j == 'Goals and Objectives':
                    query = f"""
                    The description of the subfield {j}: Specific and measurable long-term outcomes the project intends to achieve, aligned with the problem statement.

                    Instructions:
                    1. Extract and return only the content directly associated with the specific goals and objectives of the project.
                    2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                    3. Do not include any prefixes, headers, or phrases like 'the objective of the project is,' 'goals and objectives:', or similar.

                    Note:
                    - Return only the extracted goals and objectives associated with '{j}', in their original form, without rephrasing or introducing additional commentary.
                    - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                    Example:
                    Suppose the given text is 'The objective of the project is to enhance urban sustainability through smart technology integration and reduce the carbon footprint of the city by 30% within five years, and also to improve the quality of life for residents through efficient resource management.'
                    Then return: 'Enhance urban sustainability through smart technology integration. Reduce the carbon footprint of the city by 30% within five years. Improve the quality of life for residents through efficient resource management.'

                    Suppose the given text is 'The main objective of the company is to reduce the plastic quantity in the world and the deadline of the project is 12/05/2024.'
                    Then return: 'Reducing the plastic quantity in the world.'

                    Go through the instructions and examples carefully to structure the response as requested.
                    """
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        result = result.split('. ')
                    processed_output.update({j: result})
                if j == 'Expected Benefits':
                    query = f"""
                    The description of the subfield {j}: Describe the advantages the project will deliver, such as cost savings, improved efficiency, or strategic value.

                    Instructions:
                    1. Extract and return only the content directly associated with the outcomes or advantages mentioned in the text.
                    2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                    3. Do not include any prefixes, headers, or phrases like 'the expected benefits include,' 'expected benefits:', or similar.

                    Note:
                    - Return only the extracted advantages associated with '{j}', in their original form, without rephrasing or introducing additional commentary.
                    - Exclude any details about protocols, compliance requirements, dates, budgets, or unrelated information.

                    Example:
                    Suppose the given text is 'The expected benefits of this project include mnz, abc, cbz. Also includes tyz, rtz etc.'
                    Then return, 'mnz, abc, cbz. Also includes tyz, rtz etc.'

                    Suppose the given text is 'The project advantages are improved efficiency, cost savings, and strategic value.'
                    Then return, 'improved efficiency, cost savings, and strategic value.'

                    Go through the instructions and examples carefully to structure the response as requested.
                    """

                    
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})
                if j == 'Strategic Alignment':
                    query = f'''
                    The description of the subfield {j}: A description of how the project supports the organization's overall goals and priorities.

                    Instructions:
                    If the input text contains references to '{j}' or related terms, sentiments, or similar concepts, extract and return the processed content directly associated with {j}.If the given directly shares any information should be extracted Skip irrelevant content like if the content is involving the information of the other subfields like goals, problem statements etc., or lines where information is missing, incomplete, or indicates it cannot be provided.
                    Exclude any mention of unrelated phrases, such as 'no responses on the {i}------{j}' or similar.

                    Example:
                    If the input is 'Strategic Alignment of this project is mnz, abc, cbz. Also includes tyz, rtz etc', then return 'mnz, abc, cbz. Also includes tyz, rtz etc'. Don't return like 'Strategic Alignments are mnz, abc, cbz. Also includes tyz, rtz etc' or 'Strategic Alignment: mnz, abc, cbz. Also includes tyz, rtz etc'.
                    '''
                    
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})

            if i == 'Stakeholder Information':
                if j == 'Stakeholder Register':
                    query = f'''
                    If the input text contains '{j}' or related meanings, sentiments, or similar words, extract and return the relevant processed content. Skip lines where information is missing or indicates it cannot be provided. Exclude any mention of the phrase 'no responses on the {i}------{j}' or similar. Only return the extracted information without additional statements, explanations, or formatting. For example:
                    If input is 'The name of the investor is person_x. The manager will be person_A and he is also team lead for development team', then return a list of lists [['person_x','investor'], ['person_A', 'Manager',]]. Here in the provided example, it is demonstrated that a person cannot have more than one responsibility. If input is 'No responses on the field------subfield.', then return nothing.
                    '''
                    
                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Roles and Responsibilities':
                    query = f"""
                    The description of {j} is: A detailed breakdown of what each key stakeholder is responsible for in the project.

                    Instructions:
                    1. Extract and return only the relevant roles and responsibilities directly associated with '{j}' from the input text.
                    2. Include only high-level stakeholders such as project sponsors, managers, and team leads. Exclude specific roles like developers, QA engineers, or other team members unless explicitly listed as stakeholders.
                    3. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    4. Avoid introducing any additional commentary, explanations, headers, or formatting beyond the raw roles and responsibilities mentioned in the text.

                    Note:
                    - Focus only on the responsibilities of high-level stakeholders.
                    - Include the stakeholder's name or role followed by their responsibilities in plain text.
                    - If multiple stakeholders are listed, separate them clearly.
                    - Exclude unrelated details such as protocols, compliance requirements, timelines, or budgets.

                    Example:
                    Suppose the given text is "The stakeholders involved include Mr. John Doe, the project sponsor, Ms. Jane Smith, the project manager, Mr. Rajesh Kumar, the development team lead."
                    Then return, "Mr. John Doe: Project sponsor, responsible for approving budgets and resources\nMs. Jane Smith: Project manager, oversees timelines and risks\nMr. Rajesh Kumar: Development team lead, responsible for coordinating with stakeholders and managing team progress"

                    Suppose the given text is "The team includes front-end developers, back-end developers, and a UX/UI designer."
                    Then return, ""

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Carefully follow the description, instructions, and examples to provide the appropriate output for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    result = result.removeprefix('\n')
                    processed_output.update({j: result})

                if j == 'Communication Preferences':
                    query = f"""
                    The description of {j} is: Preferred communication methods (e.g., email, meetings) and the frequency of communication for each stakeholder or all stakeholders.

                    Instructions:
                    1. Extract and return only the communication preferences directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Do not include any platform names (e.g., Microsoft Teams, Zoom) unless explicitly required.
                    4. Avoid introducing '{j}' before the extracted information, additional commentary, explanations, or formatting.

                    Note:
                    - Include the stakeholder's name or role along with their preferred communication method and frequency.
                    - If multiple stakeholders or preferences are listed, separate them clearly.
                    - Exclude unrelated details such as protocols, compliance requirements, platform names, or other irrelevant information.

                    Example:
                    Suppose the given text is "Stakeholders prefer weekly meetings for updates and email for daily correspondence."
                    Then return, "Weekly meetings for updates, email for daily correspondence."

                    Suppose the given text is "The communication preference is email and two meetings per week."
                    Then return, "Email and two meetings per week."

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Carefully follow the description, instructions, and examples to provide the appropriate output for '{j}' for all stakeholders.
                    """


                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})

            if i == 'Project Scope':
                if j == 'Scope Statement':
                    query = f"""
                    The description of {j} is: A detailed description of the project’s boundaries, objectives, and deliverables, specifying what is and isn’t included.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    4. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The scope of the project includes the development of a new website, the implementation of a customer management system, and the integration with existing databases."
                    Then return, "Development of a new website, implementation of a customer management system, integration with existing databases."

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """


                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    result = result.split('. ')
                    processed_output.update({j: result})
                if j == 'Deliverables':
                    query = f"""
                    The description of {j} is: The tangible or intangible products, services, or outcomes the project will produce.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    4. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project deliverables include a fully functional website, a mobile application, and a user manual."
                    Then return, "Fully functional website, mobile application, user manual."

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})
                if j == 'Boundaries and Constraints':
                    query = f"""
                    The description of {j} is: Any limitations, such as time, budget, or resource constraints, that define the project's scope.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    4. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project is constrained by a tight timeline of 3 months, a budget of $100,000, and limited personnel."
                    Then return, "Timeline of 3 months, budget of $100,000, limited personnel."

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})

            if i == 'Requirements Documentation':
                if j == 'Functional Requirements':
                    query = f"""
                    The description of {j} is: The specific tasks, operations, or features that the project must perform (e.g., software functions, product capabilities).

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    4. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The software must support user logins, allow users to upload files, and generate reports."
                    Then return, "Support user logins, allow users to upload files, generate reports."

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})
                if j == 'Non-functional Requirements':
                    query = f"""
                    The description of {j} is: Quality attributes like performance, usability, reliability, or security that the project must meet.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    4. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The system must be able to handle 1000 users concurrently, load within 3 seconds, and be secure against data breaches."
                    Then return, "Handle 1000 users concurrently, load within 3 seconds, secure against data breaches."

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})
                if j == 'Compliance Requirements':
                    query = f"""
                    The description of {j} is: Any legal, regulatory, or organizational standards the project must adhere to.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    4. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project must comply with GDPR, HIPAA regulations, and ISO 9001 standards."
                    Then return, "GDPR, HIPAA regulations, ISO 9001 standards."

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})
                if j == 'Risk Considerations':
                    query = f"""
                    The description of {j} is: Potential risks related to the requirements and how they will be managed.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted risks and their management strategies in the following format: [['risk1', 'how it is managed'], ['risk2', 'how it is managed']].
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The risks include delays due to resource shortages, security vulnerabilities, and compliance issues. Mitigation strategies involve allocating additional resources, enhancing security protocols, and regular compliance audits."
                    Then return, [["Delays due to resource shortages", "Allocating additional resources"], ["Security vulnerabilities", "Enhancing security protocols"], ["Compliance issues", "Regular compliance audits"]]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
            if i == 'Resources':
                if j == 'Human Resources':
                    query = f"""
                    The description of {j} is: The people needed for the project, including roles, skill sets, and expertise.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the following format: ['role1', 'role2', 'role3'].
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project requires a project manager with experience in managing timelines and budgets, a software developer proficient in Python, and a designer with expertise in UI/UX."
                    Then return, ['project manager', 'software developer', 'designer']

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Physical Resources':
                    query = f"""
                    The description of {j} is: Equipment, materials, or facilities required to complete the project.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the following format: ['resource1', 'resource2'] (e.g., equipment, material, facility, etc.).
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project requires a forklift, construction materials, and office space."
                    Then return, ["Forklift", "Construction materials", "Office space"]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Budgetary Information':
                    query = f"""
                    The description of {j} is: The financial resources allocated for the project, including cost estimates and funding sources.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the following format: [['cost estimate1', 'dedicated to'], ['cost estimate2', 'dedicated to']].
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project has an estimated budget of $50,000, dedicated for marketing."
                    Then return, [["$50,000", "Marketing"]]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
            if i == 'Timeline and Deadlines':
                if j == 'High-Level Project Schedule':
                    query = f"""
                    The description of {j} is: An overview of the major phases and milestones of the project.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the following format: ['phase1', 'phase2', 'phase3']
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project has three major phases: initiation, execution, and closing. Key milestones include approval of project charter and final project review."
                    Then return, ["Initiation", "Execution", "Closing"]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Key Dates':
                    query = f"""
                    The description of {j} is: Specific important dates within the project timeline (e.g., kickoff, major review points).

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the following format: [['date1', 'event1'], ['date2', 'event2']].
                    4. Ensure the dates are sorted in chronological order and represented in the 'month day, year' format (e.g., "January 1st, 2024").
                    5. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    6. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    7. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "Key dates include project kickoff on January 1st, mid-project review on April 15th, and final project presentation on December 20th."
                    Then return, [["January 1st, 2024", "Project kickoff"], ["April 15th, 2024", "Mid-project review"], ["December 20th", "2024", "Final project presentation"]]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Deliverables':
                    query = f"""
                    The description of {j} is: Outputs expected at certain points in the timeline.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the following format: [['deliverable1', 'due date1'], ['deliverable2', 'due date2'],].
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The deliverables include project status reports due on the 15th of each month and final project report due on December 31st."
                    Then return, [["Project status report", "15th of each month"], ["Final project report", "December 31st"]]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Dependencies and Critical Path':
                    query = f"""
                    The description of {j} is: Relationships between tasks and identification of critical tasks that directly impact project completion.

                    Instructions:
                    1. Extract and strictly return only the relevant content directly associated with '{j}' from the input text.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the following format: [['task1', 'dependency1'], ['task2, dependency2']] (for task dependencies) or [['task1', 'critical path'], ['task2', 'critical path']] (for critical tasks).
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The critical path involves task A which depends on task B, and task C is dependent on task A."
                    Then return, [["Task A", "Depends on task B"], ["Task C", "Depends on task A"]]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
            if i == 'Risk Assessment':
                if j == 'Initial Risk Register':
                    query = f"""
                    The description of {j} is: A documented list of potential risks, including their likelihood and impact.

                    Instructions:
                    1. Strictly extract only the risks directly mentioned in the provided input.
                    2. If a risk has both likelihood and impact, include them. If either is missing, leave the corresponding field blank (e.g., ""). 
                    3. Return the extracted risks in the following format: [['risk1', 'likelihood1', 'impact1'], ['risk2', 'likelihood2', 'impact2']].
                    4. Do not include any information that is not related to the 'Initial Risk Register' or the risks listed in the provided input.
                    5. Skip any sections where no risks are provided (e.g., 'No responses on the {i}------{j}').
                    6. Return nothing if no risks are mentioned.
                    7. Do not add examples, inferred information, or extra commentary.
                    8. If no risk is mentioned, output nothing ("").

                    Example Suppose the given text is
                    "Risk 1: System security.; Risk 2: Resource allotment."
                    Example Then return,
                    [["System security", " ", " "], ["Resource allotment", " ", " "]]
                    """


                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})

                if j == 'Mitigation Strategies':
                    query = f"""
                    The description of {j} is: Plans and actions to reduce the likelihood or impact of identified risks.

                    Instructions:
                    1. Extract and strictly return only the mitigation strategies associated with '{j}'.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content as a list where each element represents a specific mitigation strategy. For example: ["Backup developers to address resource attrition", "Proactive communication with API providers to prevent delays"].
                    4. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    5. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    6. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "Mitigation strategies include ensuring backup developers for resource attrition and maintaining proactive communication with API providers to address potential delays."
                    Then return, ["Ensuring backup developers for resource attrition", "Maintaining proactive communication with API providers to address potential delays"]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})

                if j == 'Risk Owners':
                    query = f"""
                    The description of {j} is: Individuals responsible for monitoring and managing specific risks.

                    Instructions:
                    1. Strictly extract and return only the relevant risk owners associated with '{j}'.
                    2. If a risk owner is mentioned along with their associated risk, include them in the format: [owner, risk]. If no associated risk is mentioned directly, but a reference to previous risks (like 'these') is made, capture those risks.
                    3. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    4. Return the extracted content in the following format: [['risk1', 'owner1'], ['risk2', 'owner2']].
                    5. Do not include any information related to other subfields. Focus only on the content that is directly relevant to '{j}'.
                    6. Do not add '{j}' before the answer, or include any additional commentary, explanations, or formatting.
                    7. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "Owner 1: Sarah Williams, Risk: Data breach; Owner 2: Mark Johnson, Risk: Budget Overrun."
                    Then return, [["Data breach", "Sarah Williams"], ["Budget Overrun", "Mark Johnson"]]

                    Suppose the given text is "The project manager will monitor these risks closely."
                    Then return, [["budget overrun", "project manager"], ["resource attrition", "project manager"]]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
            if i == 'Methodology and Tools':
                if j == 'Selected Project Management Methodology':
                    query = f"""
                    The description of {j} is: The approach or framework chosen to manage the project (e.g., Agile, Waterfall).

                    Instructions:
                    1. Extract and strictly return only the relevant project management methodology mentioned in the input text for '{j}'.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content as a single string with the methodology mentioned. For example: "Agile", "Waterfall", or "Scrum".
                    4. Do not include any information related to other subfields or add extra commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project will follow the Agile Scrum methodology."
                    Then return, "Agile Scrum"

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Suppose the given text is "The methodology chosen for this project is Waterfall, focusing on sequential phases of development."
                    Then return, "Waterfall"

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                    processed_output.update({j: result})
                if j == 'Tools':
                    query = f"""
                    The description of {j} is: Software, applications, or systems used for tracking progress, collaboration, or reporting (e.g., Microsoft Project, Jira).

                    Instructions:
                    1. Extract and strictly return only the relevant tools or software mentioned in the input text for '{j}'.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content as a list of tools, e.g., ["Microsoft Project", "Jira"].
                    4. Do not include any information related to other subfields or add extra commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project will use Jira for task management and Microsoft Teams for communication."
                    Then return, ["Jira", "Microsoft Teams"]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Reporting Format and Frequency':
                    query = f"""
                    The description of {j} is: The structure and regularity of progress updates (e.g., weekly status reports, dashboards).

                    Instructions:
                    1. Extract and strictly return only the relevant reporting format and frequency mentioned in the input text for '{j}'.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content as a list, where each element includes a reporting format and its frequency, e.g., ["Weekly status reports", "Monthly dashboards"].
                    4. Do not include any information related to other subfields or add extra commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "Progress will be tracked through weekly status reports and monthly dashboards shared with stakeholders."
                    Then return, ["Weekly status reports", "Monthly dashboards"]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return, ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
            if i == 'Approval and Governance':
                if j == 'Project Charter Approval':
                    query = f"""
                    The description of {j} is: Confirmation that the project charter has been reviewed and approved by relevant authorities.

                    Instructions:
                    1. Extract and strictly return only the details about project charter approval, including who reviewed and approved it and when (if mentioned).
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the format: [["Approver1", "Date1"], ["Approver2", "Date2"]].
                    4. Do not include any information related to other subfields or add extra commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The project charter was reviewed and approved by Mr. Sandy on May 12, 2023."
                    Then return,  [["Sandy", "May 12th", "2023"]]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return,  ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'nested list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Executive Sponsor Sign-Off':
                    query = f"""
                    The description of {j} is: Formal approval and support from the project’s executive sponsor.

                    Instructions:
                    1. Extract and strictly return only the details of the executive sponsor's sign-off, including the name and date (if mentioned).
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content in the format: ["Sponsor Name", "Date"].
                    4. Do not include any information related to other subfields or add extra commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "The executive sponsor, Mr. Sandy, provided formal sign-off on December 15, 2023."
                    Then return  ["Sandy", "December 15, 2023"]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return,  ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
                if j == 'Governance Model':
                    query = f"""
                    The description of {j} is: The 'framework' for decision-making, accountability, and oversight throughout the project lifecycle.

                    Instructions:
                    1. Extract and strictly return only the relevant frameworks which acts as governance model {j} from the input. Refer the description mentioned above.
                    2. Skip any lines where information is missing, incomplete, or indicates no response is provided (e.g., 'No responses on the {i}------{j}' or similar phrases).
                    3. Return the extracted content as a list of governance elements, eg., ['Committee 1', 'committee 2'].
                    4. Do not include any information related to other subfields or add extra commentary, explanations, or formatting.
                    5. If the input is 'No responses on the field------subfield.', return nothing.

                    Example:
                    Suppose the given text is "Governance will be handled by the Project Steering Committee with monthly progress reviews and an escalation process for addressing delays."
                    Then return,  ["Project Steering Committee for oversight", "Monthly progress reviews", "Escalation process for delays"]

                    Suppose the given text is "No responses on the field------subfield."
                    Then return,  ""

                    Follow the description and instructions carefully to extract only the relevant content for '{j}'.
                    """

                    result = chain.invoke({"context": text, "question": query, 'temperature': 0})
                    result = result.content
                    if type(result) is str:
                        result = result.replace('\n', '')
                        exp = r'.+: '
                        l = regex.findall(exp, result)
                        if l:
                            result = result.replace(l[0], '')
                        res = result
                        exp = 'list'
                        human_msg = f'''input : {res}, expected format: {exp}'''
                        result = ask_gpt(sys_msg, human_msg)
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, '----------', i, '--------------', j)
                    processed_output.update({j: result})
        
        variables.charter_input.update({i: processed_output})
    print(variables.charter_input)
    return variables.charter_input

#A function which extracts the information for each fields from the given details.
def frd_extraction():
    llm = ChatOpenAI(model='gpt-3.5-turbo', temperature=0)
    prompt = PromptTemplate(
        input_variables=["context", "question"],
        template="You are a text analysis and extraction expert.\nGiven the following context:\n{context}\n"
                 "Answer as per the instructions mentioned in the question:\n{question}",
    )
    raw_text = ''
    with open(variables.project_file, 'r') as f:
        for i in f.readlines():
            raw_text += i
        text = raw_text.replace('. ', '.\n')
    chain = prompt | llm
    for i in template.template_for_frd:
        processed_output = {}
        for j in template.template_for_frd[i]:
            if i == 'Introduction':
                if j == 'Project Name':
                    processed_output.update({j: variables.charter_input['Project Overview'][j]})
                if j == 'Document Purpose':
                    text = f'''The purpose of this Functional Requirements Document (FRD) is to define the functional and non-functional requirements for {variables.charter_input['Project Overview']['Project Name']}. It serves as a key reference for stakeholders, including business analysts, developers, testers, and project managers, ensuring alignment between business needs and technical implementation.\n\nThis document outlines the system’s objectives, features, and constraints, providing a clear understanding of how the system should function. By detailing use cases, user interactions, and performance expectations, the FRD helps prevent miscommunication, minimize risks, and streamline development.\n\nAdditionally, this document establishes a structured approach to project execution, guiding teams through design, development, and testing. It ensures that [Project Name] meets business objectives, user expectations, and industry standards, leading to a successful and efficient implementation.'''
                    processed_output.update({j: text})

                if j == 'Scope Statement':
                    processed_output.update({j: variables.charter_input['Project Scope'][j]})

                if j == 'Boundaries and Constraints':
                    processed_output.update({j: variables.charter_input['Project Scope'][j]})
            if i == 'Overview':
                if j == 'Background':
                    query = '''
                    Instructions:
                    1. Summarize The whole project detail, that is problem statement, goals and objectives, scope statement and strategic alignment and return. 
                    2. If the details are not given, then return nothing.
                    3. Don't return any dummy or created information.
                    '''
                    result = chain.invoke({'context': text, 'question': query})
                    result = result.content.strip()
                    processed_output.update({j: result})

                if j == 'Goals and Objectives':
                    processed_output.update(
                        {j: variables.charter_input['Business Case and Project Objectives']['Goals and Objectives']})

                if j == 'Strategic Alignment':
                    processed_output.update(
                        {j: variables.charter_input['Business Case and Project Objectives']['Strategic Alignment']})
            if i == 'Functional Requirements':
                if j == 'Feature Description':
                    l = variables.charter_input['Requirements Documentation']['Functional Requirements']
                    sys_msg = '''
                    If the given context is providing any functional requirements separated by comma, return them in a python list datatype. If the input says it has no information or empty, then return nothing
                    '''
                    result = []
                    l = ask_gpt(sys_msg, l)
                    if l:
                        if type(l) is str:
                            l = ast.literal_eval(l)
                        for k in l:
                            query = '''
                            Describe the feature in 3-4 lines.
                            '''
                            result0 = chain.invoke({'context': text, 'question': query})
                            result.append(f'{k}:{result0.content}')
                    processed_output.update({j: result})
                if j == 'Use Cases':
                    query = f"""
                    The description of the subfield {j}: A concise description of the specific interactions between users and the system to achieve a goal, outlining the system’s expected behavior.

                     Instructions:
                    1. Extract and return only the content directly associated with the problem statement.
                    2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                    3. Do not include any prefixes, headers, or phrases like 'the problem is,' 'problem statement:', or similar.

                    Note:
                    - Return only the extracted information associated with '{j}', in its original form, without rephrasing or introducing additional commentary.
                     - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                    Example:
                    Suppose the given text is 'A user registers, logs in, resets their password, searches for products, adds items to the cart, completes checkout, an admin generates reports, a user updates their profile, a manager assigns roles, and the system sends notifications.'
                    Then return: [User Registration, Login Authentication, Password Reset, Product Search, Add to Cart, Checkout Process ,Generate Report, User Profile Update, Access Control, System Notifications]. In this example, the input is in narrative form but the output is standard words for the narration and returned in the python-list format.

                    Go through the instructions and examples carefully to structure the response as requested.
                    """
                    l1 = chain.invoke({'context': text, 'question': query}).content
                    result = []
                    if l1:
                        if type(l1) is str:
                            try:
                                l1 = ast.literal_eval(l1)
                                if l1:
                                    for k in l1:
                                        sys_msg = f'Write a single line description'
                                        m = ask_gpt(sys_msg, k)
                                        result.append(f'{k}: {m}')
                            except Exception as e:
                                print(e, f'{i} ---------------------- {j}')
                    processed_output.update({j: result})

                if j == 'User Stories':
                    query = f"""
                                The description of the subfield {j}: A concise description of system features from an end-user’s perspective, following the format: "As a [user role], I want [feature] so that [benefit]." User stories capture user needs concisely and help define functional requirements in an agile-friendly way.

                                 Instructions:
                                1. Extract and return only the content directly associated with the problem statement.
                                2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                                3. Do not include any prefixes, headers, or phrases like 'the problem is,' 'problem statement:', or similar.

                                Note:
                                - Return only the extracted information associated with '{j}', in its original form, without rephrasing or introducing additional commentary.
                                - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                                Example:
                                Suppose the given text is 'As a customer, I want to reset my password so that I can regain access to my account. As an admin, I want to generate sales reports so that I can analyze business performance. As a user, I want to receive notifications so that I stay updated about system changes.'
                                Then return: ["A customer wants to reset their password to regain access to their account.", "An admin needs to generate sales reports to analyze business performance.", "A user requires notifications to stay updated about system changes."] In this example, the input is in direct speech, but the output is structured in indirect speech and returned in the python-list format.

                                Go through the instructions and examples carefully to structure the response as requested.
                                """
                    result = chain.invoke({'context': text, 'question': query}).content
                    if type(result) is str:
                        try:
                            result = ast.literal_eval(result)
                        except Exception as e:
                            print(e, f'{i} ---------------------- {j}')
                    processed_output.update({j: result})

                if j == 'Input & Output Requirements':
                    query = f"""
                                The description of the subfield {j}: A detailed specification of the data inputs required by the system and the corresponding outputs generated. Inputs define the data the system receives, while outputs describe the expected results after processing.

                                 Instructions:
                                1. Extract and return only the content directly associated with the problem statement.
                                2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                                3. Do not include any prefixes, headers, or phrases like 'the problem is,' 'problem statement:', or similar.

                                Note:
                                - Return only the extracted information associated with '{j}', in its original form, without rephrasing or introducing additional commentary.
                                - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                                Example:
                                Suppose the given text is 'The system accepts a username and password as input and returns authentication success or failure. A user provides search keywords, and the system outputs a list of relevant products. A payment form collects credit card details, and the system confirms the transaction status. A file upload feature takes a document as input and generates a downloadable processed file.'
                                Then return: [["Username and Password", "Authentication Success/Failure"], ["Search Keywords", "Relevant Product List"], ["Credit Card Details", "Transaction Status"], ["Document Upload", "Processed Downloadable File"]]. In this example, the input is in narrative form, but the output is structured as a nested list where each sublist contains an input and its corresponding output.

                                Go through the instructions and examples carefully to structure the response as requested.
                                """
                    result = chain.invoke({'context': text, 'question': query}).content
                    if result:
                        if type(result) is str:
                            try:
                                result = ast.literal_eval(result)
                            except Exception as e:
                                print(e, f'{i} ---------------------- {j}')
                        processed_output.update({j: result})
            if i == 'Non Functional Requirements':
                if j == 'Performance Requirements':
                    query = f"""
                                The description of the subfield {j}: Defines the expected system performance, including response times, throughput, scalability, and resource utilization under specific conditions.

                                 Instructions:
                                1. Extract and return only the content directly associated with the problem statement.
                                2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                                3. Do not include any prefixes, headers, or phrases like 'the problem is,' 'problem statement:', or similar.

                                Note:
                                - Return only the extracted information associated with '{j}', in its original form, without rephrasing or introducing additional commentary.
                                - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                                Example:
                                Suppose the given text is 'The system should handle up to 10,000 concurrent users with a response time of under 2 seconds. The database should process 5,000 transactions per second. The API should have 99.99% uptime. The system should scale horizontally to support increased traffic without performance degradation.'
                                Then return: ["Handle up to 10,000 concurrent users with a response time under 2 seconds.", "Process 5,000 transactions per second.", "Ensure API uptime of 99.99%.", "Scale horizontally to support increased traffic without performance degradation."]. In this example, the input is in narrative form, but the output is structured as a list of clear, measurable performance requirements.

                                Go through the instructions and examples carefully to structure the response as requested.
                                """
                    result = chain.invoke({'context': text, 'question': query}).content
                    if type(result) is str:
                        try:
                            result = ast.literal_eval(result)
                        except Exception as e:
                            print(e, f'{i} ---------------------- {j}')
                    processed_output.update({j: result})
                if j == 'Security Requirements':
                    query = f"""
                                The description of the subfield {j}: Specifies the security measures, controls, and protocols required to protect data, users, and system integrity against unauthorized access, breaches, and other threats.

                                 Instructions:
                                1. Extract and return only the content directly associated with the problem statement.
                                2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                                3. Do not include any prefixes, headers, or phrases like 'the problem is,' 'problem statement:', or similar.

                                Note:
                                - Return only the extracted information associated with '{j}', in its original form, without rephrasing or introducing additional commentary.
                                - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                                Example:
                                Suppose the given text is 'The system must enforce multi-factor authentication for all user logins. Data should be encrypted at rest and in transit using AES-256 encryption. Role-based access control should restrict user permissions. The system should automatically log out inactive users after 15 minutes of inactivity.'
                                Then return: ["Enforce multi-factor authentication for user logins.", "Encrypt data at rest and in transit using AES-256.", "Implement role-based access control to restrict user permissions.", "Auto log out inactive users after 15 minutes."]. In this example, the input is in narrative form, but the output is structured as a list of specific security requirements.

                                Go through the instructions and examples carefully to structure the response as requested.
                                """
                    result = chain.invoke({'context': text, 'question': query}).content
                    if type(result) is str:
                        try:
                            result = ast.literal_eval(result)
                        except Exception as e:
                            print(e, f'{i} ---------------------- {j}')
                    processed_output.update({j: result})
                if j == 'Scalability':
                    query = f"""
                                The description of the subfield {j}: Defines the system’s ability to handle growth in users, data volume, and workload without performance degradation, including horizontal and vertical scaling strategies.

                                 Instructions:
                                1. Extract and return only the content directly associated with the problem statement.
                                2. Skip any irrelevant lines, incomplete information, or references indicating that no information is available (e.g., 'no responses on the {i}------{j}' or similar).
                                3. Do not include any prefixes, headers, or phrases like 'the problem is,' 'problem statement:', or similar.

                                Note:
                                - Return only the extracted information associated with '{j}', in its original form, without rephrasing or introducing additional commentary.
                                - Exclude any details about project names, dates, compliance requirements, budgets, or unrelated information.

                                Example:
                                Suppose the given text is 'The system should support auto-scaling to handle traffic spikes efficiently. It must be capable of distributing workload across multiple servers. Database sharding should be implemented to manage large datasets. Caching mechanisms should reduce response times under high loads.'
                                Then return: ["Support auto-scaling to handle traffic spikes.", "Distribute workload across multiple servers.", "Implement database sharding for large datasets.", "Use caching mechanisms to reduce response times under high loads."]. In this example, the input is in narrative form, but the output is structured as a list of specific scalability requirements.

                                Go through the instructions and examples carefully to structure the response as requested.
                                """
                    result = chain.invoke({'context': text, 'question': query}).content
                    if type(result) is str:
                        try:
                            result = ast.literal_eval(result)
                        except Exception as e:
                            print(e, f'{i} ---------------------- {j}')
                    processed_output.update({j: result})

            if i == 'Risks & Mitigations':
                if j == 'Initial Risk Register':
                    result = variables.charter_input['Risk Assessment'][j]
                    processed_output.update({j: result})
                if j == 'Risks & Mitigations':
                    result0 = variables.charter_input['Requirements Documentation']['Risk Considerations']
                    result = []
                    if result0:
                        if type(result0) is str:
                            try:
                                result0 = ast.literal_eval(result0)
                            except Exception as e:
                                print(e, f'{i} ---------- {j}')
                        if type(result0) is list:
                            for x in result0:
                                sys_msg = f"The risk in the project is {x[0]} and the mitigation strategy is {x[1]}. Now explain these two, how the mitigation helps is resolving the risk in 3-4 short lines. Return the explained answer"
                                result.append(ask_gpt(sys_msg, str(x)))
                    processed_output.update({j: result})

            if i == 'Dependencies':
                if j == 'Dependencies and Critical Path':
                    result = variables.charter_input['Timeline and Deadlines']['Dependencies and Critical Path']
                    processed_output.update({j: result})

        variables.frd_input.update({i: processed_output})
    return variables.frd_input
#class containing functionalities to generate the charter and frd
class generate_charter:
    charter_input = {}
    def __init__(self, charter_input:dict):
        self.charter_input = charter_input

    def add_heading(self,doc, text, color):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        paragraph._p.get_or_add_pPr().append(shading_elm)

    def add_subfield(self,doc, text):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    def add_content(self,doc, text):
        if isinstance(text, list):
            for item in text:
                if isinstance(item, list):
                    # Handle nested lists (e.g., for dependencies)
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(" - ".join(str(x) for x in item))
                else:
                    paragraph = doc.add_paragraph(str(item), style="List Bullet")
        else:
            lines = text.split("\n")
            for line in lines:
                paragraph = doc.add_paragraph(line.strip(), style="List Bullet" if len(lines) > 1 else None)

    def set_cell_background(self,cell, color):
        """Apply background color to a table cell."""
        cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
        shd = cell._element.tcPr.find(qn('w:shd'))
        shd.set(qn('w:fill'), color)  # Set the fill color

    def add_table(self,doc, data, field):

        head =[]
        if field == 'Stakeholder Register':
            head = ["Name", "Role"]
        if field == 'Risk Considerations':
            head = ["Risk", "Mitigation"]
        if field == 'Budgetary Information':
            head = ["Budget", "Dedicated to"]
        if field == 'Key Dates':
            head = ["Milestone", "Date"]
        if field == 'Deliverables':
            head = ["Deliverables", "Date"]
        if field == 'Dependencies and Critical Path':
            head = ["Process", "Depends on/ Critical path"]
        if field == 'Initial Risk Register':
            head = ["Risk", "Likelihood", "Impact"]
        if field == 'Risk Owners':
            head = ["Risk", "Owner"]
        if field == 'Project Charter Approval':
            head = ["Approver", "Date"]

        table = doc.add_table(rows=1, cols=len(head))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells

        for i, column in enumerate(head):
            hdr_cells[i].text = str(column)
            self.set_cell_background(hdr_cells[i], "D9E1F2")  # Light Blue Header Background

        if not data:
            return table

        for row in data[0:]:
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

        return table

    def create_project_charter(self, file_path):
        project_data = self.charter_input
        doc = Document()

        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run("Project Charter")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Iterate through the data dictionary and populate the Word document
        for section_title, content in project_data.items():
            self.add_heading(doc, section_title, "CAE4F1")

            if isinstance(content, dict):
                for subfield, value in content.items():
                    self.add_subfield(doc, subfield)
                    if isinstance(value, list) and value and isinstance(value[0], list):
                        self.add_table(doc, value, subfield)
                    else:
                        self.add_content(doc, value)
            else:
                self.add_content(doc, content)

            # Add spacing between sections
            doc.add_paragraph()

        doc.save(file_path)

class generate_frd:
    frd_input = {}
    def __init__(self, frd_input:dict):
        self.frd_input = frd_input

    def add_heading(self,doc, text, color):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        paragraph._p.get_or_add_pPr().append(shading_elm)

    def add_subfield(self,doc, text):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    def add_content(self,doc, text):
        if isinstance(text, list):
            for item in text:
                if isinstance(item, list):
                    # Handle nested lists (e.g., for dependencies)
                    paragraph = doc.add_paragraph()
                    paragraph.add_run(" - ".join(str(x) for x in item))
                else:
                    paragraph = doc.add_paragraph(str(item), style="List Bullet")
        else:
            lines = text.split("\n")
            for line in lines:
                paragraph = doc.add_paragraph(line.strip(), style="List Bullet" if len(lines) > 1 else None)

    def set_cell_background(self,cell, color):
        """Apply background color to a table cell."""
        cell._element.get_or_add_tcPr().append(OxmlElement('w:shd'))
        shd = cell._element.tcPr.find(qn('w:shd'))
        shd.set(qn('w:fill'), color)  # Set the fill color

    def add_table(self,doc, data, field):
        try:
            head = []
            if field == 'Initial Risk Register':
                head = ["Risk", "Likelihood", "Impact"]
            if field == 'Dependencies and Critical Path':
                head = ["Process", "Depends on/ Critical path"]
            if field == 'Input & Output Requirements':
                head = ["Input", "Output"]

            table = doc.add_table(rows=1, cols=len(head))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells

            for i, column in enumerate(head):
                hdr_cells[i].text = str(column)
                self.set_cell_background(hdr_cells[i], "D9E1F2")  # Light Blue Header Background

            if not data:
                return table

            for row in data[0:]:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = str(cell)

            return table
        except Exception as e:
            print(e, f"\t\t\t\t\t{field}")

    def create_frd(self, file_path):
        project_data = self.frd_input
        doc = Document()

        # Add title
        title = doc.add_paragraph()
        title_run = title.add_run("Functional Requirements Documentation (FRD)")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Iterate through the data dictionary and populate the Word document
        for section_title, content in project_data.items():
            self.add_heading(doc, section_title, "CAE4F1")

            if isinstance(content, dict):
                for subfield, value in content.items():
                    self.add_subfield(doc, subfield)
                    if isinstance(value, list) and value and isinstance(value[0], list):
                        self.add_table(doc, value, subfield)
                    else:
                        self.add_content(doc, value)
            else:
                self.add_content(doc, content)

            # Add spacing between sections
            doc.add_paragraph()

        doc.save(file_path)


